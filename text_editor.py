import tkinter
from tkinter import *
from tkinter import filedialog
import textract
import subprocess
from docx import Document

# В будущем испольщовать rope структуры данных вместо готовых компонентов
# идея сделать текстовый редакор с встроенным AI. Короче хочу сделать дописывание текста как в notion.
# добавить ctrl+z
class Editor:
    def __init__(self, window):
        self.window = window
        self.window.title("My text editor") #название окна

        self.text_area = tkinter.Text(self.window)
        self.text_area.pack(expand=True, fill='both')

        self.menubar = tkinter.Menu(self.window)
        file_menu = tkinter.Menu(self.menubar, tearoff=0)
        file_menu.add_command(label="New", command=self.new_file)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="save", command=self.save_file)

        self.menubar.add_cascade(label="File", menu=file_menu)

        self.window.config(menu=self.menubar)

        self.current_file_path = None # Атрибут для хранения пути к файлу

        self.shortcuts()
        
        # Вызываем метод для уведомлений
        self.status_bar = Label(self.window, text="", bd=1, relief=SUNKEN, anchor=W)
        self.status_bar.pack(side=BOTTOM, fill=X)

    def new_file(self):
        pass

    def open_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt"), ("Word documents", "*.doc;*.docx")])
        if file_path:
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                # para это просто переменная. Она не инициализируется явно
                # join объеденяет элементы
                # \n разделитель. В данном случае каждый новый элемент с новой строки
                content = '\n'.join([para.text for para in doc.paragraphs])
            elif file_path.endswith('.doc'):
                # unoconv для ковертации .doc в текст
                content = subprocess.check_output(['unoconv', '-f', 'text', file_path]).decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as file: #открываем файл в режиме чтения
                    content = file.read()

            self.text_area.delete(1.0, tkinter.END) # tkinter.END константа либы тк
            self.text_area.insert(tkinter.END, content)
# TODO Кнопка "сохранить как"
    def save_file(self, event=None):
        if self.current_file_path:
            text = self.text_area.get("1.0", tkinter.END)
            if self.current_file_path.endswith('.docx'):
                doc = Document() # создаём экземпляр класса Document
                doc.add_paragraph(text) # метод класса Document
                doc.save(self.current_file_path)
                self.update_status("Saved")
            else:
                with open(self.current_file_path, 'w', encoding='utf-8') as file:
                    file.write(text)
                self.update_status("Saved")
        else:
            file_path = filedialog.asksaveasfilename(filetypes=[("Text files", "*.txt"), ("Word documents", ".docx")])
            if file_path:
                text = self.text_area.get("1.0", tkinter.END)
                if file_path.endswith('.docx'):
                    doc = Document()
                    doc.add_paragraph(text)
                    doc.save(file_path)
                    self.current_file_path = file_path
                else:
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(text)
                        self.current_file_path = file_path # Сохраняем путь к файлу

    def update_status(self,message):
        # обновляем текст на панели состояния с помощью переданного сообщения
        self.status_bar.config(text=message)
        # Изменяем текст 5 секунд
        self.window.after(5000, lambda: self.status_bar.config(text=""))

    def shortcuts(self):
        self.text_area.bind("<Control-s>",self.save_file)

# TODO настройка отступов и интервалов у пользователя + интерфейс линейка для настройки красной строки
window = tkinter.Tk() # взяли из тк класс и создали его экземпляр
Editor(window) # создаём экземпляр класса эдитор и передаём ему виндов
window.mainloop()
